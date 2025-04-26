import os
import csv
from docx import Document

# python -m NCR_old_prices_get_csv


def write_last_three_cols_to_csv(docx_path: str, word_name: str):
    try:
        # Open the Word document
        doc = Document(docx_path)

        # Check if there are any tables in the document
        if not doc.tables:
            print("Error: No tables found in the document.")
            return

        # Get the first table
        table = doc.tables[0]

        cols = len(table.columns)

        # Check if the table has at least two columns
        if cols < 3:
            print("Error: The table has less than two columns.")
            return

        # Iterate through the rows and print the last two columns
        yrs = ["2019", "2020", "2021", "2022", "2023", "2024", "2025"]
        year = ""
        for yr in yrs:
            if yr in word_name:
                year = yr
                break
        output_csv_path: str = f"doe_csvs\\old_format\\{year}\\{word_name.replace(".docx", ".csv")}"
        with open(output_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=["City", "Product", "Min Price", "Max Price", "Common Price"])
            writer.writeheader()
            for row in table.rows:
                cell_0 = clean_text(row.cells[0].text)
                prod_cell = clean_text(row.cells[1].text)
                cell_1 = clean_text(row.cells[cols - 3].text)
                cell_2 = clean_text(row.cells[cols - 2].text)
                cell_3 = clean_text(row.cells[cols - 1].text)
                if cell_1 == "OVERALL RANGE" or cell_1 == "NATIONAL CAPITAL REGION":
                    continue
                writer.writerow({'City': cell_0, "Product": prod_cell, 'Min Price': cell_1, 'Max Price': cell_2, 'Common Price': cell_3})
        print(f"Data successfully written to {output_csv_path}")

    except FileNotFoundError:
        print(f"Error: File not found at {docx_path}")
    except Exception as e:
        print(f"An error occurred: {e}")


def clean_text(text: str) -> str:
    text = text.replace('\t', '')
    text = text.replace('-', '')
    if text == "":
        return "#N/A"
    return text


if __name__ == "__main__":
    raw_data = "doe_word"
    for folder in os.listdir(raw_data):
        year_folder = os.path.join(raw_data, folder)
        if not ("2021" in year_folder):
            continue
        for filename in os.listdir(year_folder):
            word_file =  os.path.join(year_folder, filename)
            if not ("ncr" in filename): ## only convert NCR data
                continue
            write_last_three_cols_to_csv(os.path.join(os.getcwd(), word_file), filename)