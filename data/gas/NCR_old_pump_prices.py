import os
import csv
from docx import Document
import comtypes.client

def open_pdf_in_word(pdf_path: str, docx_path: str, pdf_name: str):
    """
    Opens a PDF file directly in Microsoft Word.  This relies on Word's ability to
    open PDF files, which may not always produce perfect results.

    Args:
        pdf_path (str): The path to the PDF file.
    """
    try:
        # Check if the PDF file exists
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Error: PDF file not found at {pdf_path}")

        # Open the PDF file using Microsoft Word
        #subprocess.run(['start', 'winword', pdf_path], shell=True)
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = True  # Make Word visible to see what's happening
        doc = word.Documents.Open(os.path.join(os.getcwd(), pdf_path))  # Open the PDF
        # time.sleep(2)  # Give Word a moment to process, may need adjustment
        doc.SaveAs(os.path.join(os.getcwd(), docx_path), FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()

        ## now print the last 2 columns
        write_last_three_cols_to_csv(os.path.join(os.getcwd(), docx_path), pdf_name)
        print(f"Successfully saved to {docx_path} using COM automation.")

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An error occurred: {e}")


def write_last_three_cols_to_csv(docx_path: str, pdf_name: str):
    try:
        # Open the Word document
        doc = Document(docx_path)

        # Check if there are any tables in the document
        if not doc.tables:
            print("Error: No tables found in the document.")
            return

        # Get the first table
        table = doc.tables[0]

        cols = len(table.columns)

        # Check if the table has at least two columns
        if cols < 3:
            print("Error: The table has less than two columns.")
            return

        # Iterate through the rows and print the last two columns
        yrs = ["2019", "2020", "2021", "2022", "2023", "2024", "2025"]
        year = ""
        for yr in yrs:
            if yr in pdf_name:
                year = yr
                break
        output_csv_path: str = f"doe_csvs\\old_format\\{year}\\{pdf_name.replace(".pdf", ".csv")}"
        with open(output_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=["City", "Min Price", "Max Price", "Common Price"])
            writer.writeheader()
            for row in table.rows:
                cell_0 = clean_text(row.cells[0].text)
                cell_1 = clean_text(row.cells[cols - 3].text)
                cell_2 = clean_text(row.cells[cols - 2].text)
                cell_3 = clean_text(row.cells[cols - 1].text)
                if cell_1 == "OVERALL RANGE" or cell_1 == "NATIONAL CAPITAL REGION":
                    continue
                writer.writerow({'City': cell_0, 'Min Price': cell_1, 'Max Price': cell_2, 'Common Price': cell_3})
        print(f"Data successfully written to {output_csv_path}")

    except FileNotFoundError:
        print(f"Error: File not found at {docx_path}")
    except Exception as e:
        print(f"An error occurred: {e}")



def clean_text(text: str) -> str:
    text = text.replace('\t', '')
    text = text.replace('-', '')
    if text == "":
        return "#N/A"
    return text


if __name__ == "__main__":
    raw_data = "doe_pdfs\\old_format"
    for folder in os.listdir(raw_data):
        year_folder = os.path.join(raw_data, folder)
        if not ("2020" in year_folder): ## choose specific year
            continue
        for filename in os.listdir(year_folder):
            pdf_file =  os.path.join(year_folder, filename)
            output_path = f"doe_word\\{folder}\\{filename.replace(".pdf", ".docx")}"
            if not ("ncr" in filename): ## only convert NCR data
                continue
            if os.path.exists(output_path):
                continue
            if os.path.exists(pdf_file):
                open_pdf_in_word(pdf_file, output_path, filename)
            else:
                print(f"Error: PDF file '{pdf_file}' not found.")